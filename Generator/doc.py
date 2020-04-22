import docx
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import argparse

ap = argparse.ArgumentParser()
ap.add_argument("-f", "--font", required=True,
	            help="Font Type")
ap.add_argument("-s", "--size", required=True,
                help="Font Size")
ap.add_argument("-r", "--repetition", help="Number of char in one line")
ap.add_argument("-d", "--dot", help="Space")

args = vars(ap.parse_args())
font_arg = args["font"]
size_arg = int(args["size"])
if args["repetition"] != None:
    repetition_arg = int(args["repetition"])
else:
    repetition_arg = ''
if args["dot"] != None:
    dot_arg = int(args["dot"])
else:
    dot_arg = ''

original_list = ['HAMZA',
                 'ALEEF with HAMZA above',
                 'ALEEF with HAMZA below',
                 'ALEEF',
                 'BEH',
                 'TEH MARBUTA',
                 'TEH',
                 'THEH',
                 'JEEM',
                 'HAH',
                 'KHAH',
                 'DAL',
                 'THAL',
                 'REH',
                 'ZAIN',
                 'SEEN',
                 'SHEEN',
                 'SAD',
                 'DAD',
                 'TAH',
                 'ZAH',
                 'AIN',
                 'GHAIN',
                 'FEH',
                 'QAF',
                 'KAF',
                 'LAM',
                 'MEEM',
                 'NOON',
                 'HEH',
                 'WAW',
                 'YEH',
                 ]
original = [u'\u0621',  # 0
            u'\u0623',  # 1
            u'\u0625',  # 2
            u'\u0627',  # 3
            u'\u0628',  # 4
            u'\u0629',  # 5
            u'\u062A',  # 6
            u'\u062B',  # 7
            u'\u062C',  # 8
            u'\u062D',  # 9
            u'\u062E',  # 10
            u'\u062F',  # 11
            u'\u0630',  # 12
            u'\u0631',  # 13
            u'\u0632',  # 14
            u'\u0633',  # 15
            u'\u0634',  # 16
            u'\u0635',  # 17
            u'\u0636',  # 18
            u'\u0637',  # 19
            u'\u0638',  # 20
            u'\u0639',  # 21
            u'\u063A',  # 22
            u'\u0641',  # 23
            u'\u0642',  # 24
            u'\u0643',  # 25
            u'\u0644',  # 26
            u'\u0645',  # 27
            u'\u0646',  # 28
            u'\u0647',  # 29
            u'\u0648',  # 30
            u'\u064A',  # 31
            ]
forms_b = [ # u'\uFE80',  # ALEEF isolated
            u'\u0621',
            # u'\uFE83',
            # u'\uFE87',
            # u'\uFE8D',
            u'\u0623',
            u'\u0625',
            u'\u0627',
            u'\uFE84',  # _end
            u'\uFE88',
            u'\uFE8E',
            # u'\uFE8F',  # BEH isolated
            u'\u0628',
            u'\uFE91',  # _begin
            u'\uFE92',  # _middle
            u'\uFE90',  # _end
            # u'\uFE95',  # TEH isolated
            u'\u062A',
            # u'\uFE93',
            u'\u0629',
            u'\uFE97',  # _begin
            u'\uFE98',  # _middle
            u'\uFE94',  # _end
            u'\uFE96',
            # u'\uFE99',  # THEH isolated
            u'\u062B',
            u'\uFE9B',  # _begin
            u'\uFE9C',  # _middle
            u'\uFE9A',  # _end
            # u'\uFE9D',  # JEEM isolated
            u'\u062C',
            u'\uFE9F',  # _begin
            u'\uFEA0',  # _middle
            u'\uFE9E',  # _end
            # u'\uFEA1',  # HAH isolated
            u'\u062D',
            u'\uFEA3',  # _begin
            u'\uFEA4',  # _middle
            u'\uFEA2',  # _end
            # u'\uFEA5',  # KHAH isolated
            u'\u062E',
            u'\uFEA7',  # _begin
            u'\uFEA8',  # _middle
            u'\uFEA6',  # _end
            # u'\uFEA9',  # DAL isolated
            u'\u062F',
            u'\uFEAA',  # _end
            # u'\uFEAB',  # THAL isolated
            u'\u0630',
            u'\uFEAC',  # _end
            # u'\uFEAD',  # REH isolated
            u'\u0631',
            u'\uFEAE',  # _end
            # u'\uFEAF',  # ZAIN isolated
            u'\u0632',
            u'\uFEB0',  # _end
            # u'\uFEB1',  # SEEN isolated
            u'\u0633',
            u'\uFEB3',  # _begin
            u'\uFEB4',  # _middle
            u'\uFEB2',  # _end
            # u'\uFEB5',  # SHEEN isolated
            u'\u0634',
            u'\uFEB7',  # _begin
            u'\uFEB8',  # _middle
            u'\uFEB6',  # _end
            # u'\uFEA9',  # SAD isolated
            u'\u0635',
            u'\uFEBB',  # _begin
            u'\uFEBC',  # _middle
            u'\uFEBA',  # _end
            # u'\uFEBD',  # DAD isolated
            u'\u0636',
            u'\uFEBF',  # _begin
            u'\uFEC0',  # _middle
            u'\uFEBE',  # _end
            # u'\uFEC1',  # TAH isolated
            u'\u0637',
            u'\uFEC3',  # _begin
            u'\uFEC4',  # _middle
            u'\uFEC2',  # _end
            # u'\uFEC5',  # ZAH isolated
            u'\u0638',
            u'\uFEC7',  # _begin
            u'\uFEC8',  # _middle
            u'\uFEC6',  # _end
            # u'\uFEC9',  # AIN isolated
            u'\u0639',
            u'\uFECB',  # _begin
            u'\uFECC',  # _middle
            u'\uFECA',  # _end
            # u'\uFECD',  # GHAIN isolated
            u'\u063A',
            u'\uFECF',  # _begin
            u'\uFED0',  # _middle
            u'\uFECE',  # _end
            # u'\uFED1',  # FEH isolated
            u'\u0641',
            u'\uFED3',  # _begin
            u'\uFED4',  # _middle
            u'\uFED2',  # _end
            # u'\uFED5',  # QAF isolated
            u'\u0642',
            u'\uFED7',  # _begin
            u'\uFED8',  # _middle
            u'\uFED6',  # _end
            # u'\uFED9',  # KAF isolated
            u'\u0643',
            u'\uFEDB',  # _begin
            u'\uFEDC',  # _middle
            u'\uFEDA',  # _end
            # u'\uFEDD',  # LAM isolated
            u'\u0644',
            u'\uFEFB',
            u'\uFEDF',  # _begin
            u'\uFEE0',  # _middle
            u'\uFEDE',  # _end
            u'\uFEFC',
            # u'\uFEE1',  # MEEM isolated
            u'\u0645',
            u'\uFEE3',  # _begin
            u'\uFEE4',  # _middle
            u'\uFEE2',  # _end
            # u'\uFEE5',  # NOON isolated
            u'\u0646',
            u'\uFEE7',  # _begin
            u'\uFEE8',  # _middle
            u'\uFEE6',  # _end
            # u'\uFEE9',  # HEH isolated
            u'\u0647',
            u'\uFEEB',  # _begin
            u'\uFEEC',  # _middle
            u'\uFEEA',  # _end
            # u'\uFEED',  # WAW isolated
            u'\u0648',
            u'\uFEEE',  # _end
            # u'\uFEF1',  # YEH isolated
            u'\u0649',
            u'\uFEF3',  # _begin
            u'\uFEF4',  # _middle
            u'\uFEF2'  # _end
            
    ]
tashkhil_ = ['fathatan', 'dhammahtan', 'kasratan', 'fatha',
            'damma', 'kasra', 'shadda', 'sukun', 'sukun_c']
tashkhil_uni = [u'\u064B',u'\u064C',u'\u064D',u'\u064E',
            u'\u064F',u'\u0650',u'\u0651',u'\u0652',u'\u06E1']


def form_A_dict(tashkhil):
    if tashkhil:
        v = tashkhil_uni[3]
    else:
        v = ''
    form_A_dict ={}
    form_A_dict['ALEEF'] = [
        original[25] + v + original[3] + v ,
        original[27] + v + original[3] + v ,
        original[3] + v + original[26] + v + original[26] + v + original[29] + v ,
        original[3] + v + original[25] + v + original[4] + v + original[13] + v
    ]
    form_A_dict['BEH'] = [
        original[4] + v + original[8] + v ,  # Isolated
        original[4] + v + original[9] + v ,
        original[4] + v + original[10] + v ,
        original[4] + v + original[27] + v ,
        original[4] + v + original[31] + v ,
        original[4] + v + original[4] + v + original[31] + v , # Final + same front
        original[4] + v + original[4] + v + original[13] + v ,
        original[4] + v + original[4] + v + original[14] + v ,
        original[4] + v + original[4] + v + original[27] + v ,
        original[4] + v + original[4] + v + original[28] + v ,
        original[4] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[4] + v + original[27] + v + original[3] + v ,
        original[4] + v + original[29] + v + original[3] + v ,
        original[4] + v + original[8] + v + original[3] + v ,
        original[4] + v + original[10] + v + original[3] + v ,
        original[4] + v + original[4] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[4] + v + original[4] + v + original[29] + v + original[3] + v ,
        original[4] + v + original[4] + v + original[9] + v + original[31] + v , # final
        original[4] + v + original[4] + v + original[10] + v + original[31] + v , 
        original[3] + v + original[25] + v + original[4] + v + original[13] + v
    ]
    form_A_dict['TEH'] = [
        original[6] + v + original[8] + v ,  # Isolated
        original[6] + v + original[9] + v ,
        original[6] + v + original[10] + v ,
        original[6] + v + original[27] + v ,
        original[6] + v + original[31] + v ,
        
        original[6] + v + original[6] + v + original[13] + v , # Final + same front
        original[6] + v + original[6] + v + original[14] + v ,
        original[6] + v + original[6] + v + original[27] + v ,
        original[6] + v + original[6] + v + original[28] + v ,
        original[6] + v + original[6] + v + original[31] + v ,
        original[6] + v + original[6] + v + original[9] + v + original[8] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[8] + v + original[31] + v ,
        original[6] + v + original[6] + v + original[10] + v + original[31] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[27] + v + original[31] + v ,
        
        original[6] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[6] + v + original[9] + v + original[3] + v ,
        original[6] + v + original[10] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[3] + v ,
        original[6] + v + original[29] + v + original[3] + v ,
        original[6] + v + original[8] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[9] + v + original[8] + v + original[3] + v ,
        original[6] + v + original[9] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[10] + v + original[27] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[8] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[27] + v + original[9] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[10] + v + original[3] + v , # medial + same front + alif end
        
        original[6] + v + original[6] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[29] + v + original[3] + v 
    ]
    form_A_dict['THEH'] = [
        original[7] + v + original[8] + v ,  # Isolated
        original[7] + v + original[27] + v ,
        original[7] + v + original[31] + v ,
        
        original[7] + v + original[7] + v + original[13] + v , # Final + same front
        original[7] + v + original[7] + v + original[14] + v ,
        original[7] + v + original[7] + v + original[27] + v ,
        original[7] + v + original[7] + v + original[28] + v ,
        original[7] + v + original[7] + v + original[31] + v ,

        original[7] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[7] + v + original[7] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[7] + v + original[7] + v + original[29] + v + original[3] + v 
    ]
    form_A_dict['JEEM'] = [
        original[8] + v + original[9] + v ,  # Isolated
        original[8] + v + original[27] + v ,
        original[8] + v + original[31] + v ,
        
        original[8] + v + original[8] + v + original[31] + v , # Final + same front
        original[8] + v + original[8] + v + original[27] + v + original[9] + v ,
        original[8] + v + original[8] + v + original[9] + v + original[31] + v,
        original[8] + v + original[8] + v + original[27] + v + original[31] + v,

        original[8] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[8] + v + original[8] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[8] + v + original[8] + v + original[29] + v + original[3] + v 
    ]
    form_A_dict['HAH'] = [
        original[9] + v + original[8] + v ,  # Isolated
        original[9] + v + original[27] + v ,
        original[9] + v + original[31] + v ,
        
        original[9] + v + original[9] + v + original[31] + v , # Final + same front
        original[9] + v + original[9] + v + original[27] + v + original[31] + v ,
        original[9] + v + original[9] + v + original[8] + v + original[31] + v,

        original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[9] + v + original[27] + v + original[3] + v  # Inital + alif end

    ]
    form_A_dict['KHAH'] = [
        original[10] + v + original[9] + v ,  # Isolated
        original[10] + v + original[27] + v ,
        original[10] + v + original[8] + v ,
        original[10] + v + original[31] + v ,

        original[10] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[10] + v + original[31] + v + original[3] + v  # Inital + alif end

    ]
    form_A_dict['SEEN'] = [
        original[15] + v + original[8] + v ,  # Isolated
        original[15] + v + original[9] + v ,  # Isolated
        original[15] + v + original[10] + v ,  # Isolated
        original[15] + v + original[13] + v ,  # Isolated
        original[15] + v + original[27] + v ,
        original[15] + v + original[31] + v ,
        
        original[15] + v + original[15] + v + original[13] + v , # Final + same front
        original[15] + v + original[15] + v + original[31] + v ,
        original[15] + v + original[15] + v + original[27] + v + original[9] + v ,
        original[15] + v + original[15] + v + original[27] + v + original[27] + v ,
        original[15] + v + original[15] + v + original[10] + v + original[31] + v ,

        original[15] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[15] + v + original[15] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[15] + v + original[15] + v + original[29] + v + original[3] + v ,
        original[15] + v + original[15] + v + original[8] + v + original[3] + v ,
        original[15] + v + original[15] + v + original[9] + v + original[3] + v ,
        original[15] + v + original[15] + v + original[10] + v + original[3] + v

    ]
    form_A_dict['SHEEN'] = [
        original[16] + v + original[8] + v ,  # Isolated
        original[16] + v + original[9] + v ,  # Isolated
        original[16] + v + original[10] + v ,  # Isolated
        original[16] + v + original[13] + v ,  # Isolated
        original[16] + v + original[27] + v ,
        original[16] + v + original[31] + v ,
        
        original[16] + v + original[16] + v + original[8] + v , # Final + same front
        original[16] + v + original[16] + v + original[9] + v ,
        original[16] + v + original[16] + v + original[10] + v, # Final + same front
        original[16] + v + original[16] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[13] + v , # Final + same front
        original[16] + v + original[16] + v + original[31] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[8] + v + original[31] + v ,
        original[16] + v + original[16] + v + original[27] + v + original[10] + v ,
        original[16] + v + original[16] + v + original[27] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[31] + v ,

        original[16] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[16] + v + original[16] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[16] + v + original[16] + v + original[29] + v + original[3] + v ,
        original[16] + v + original[16] + v + original[8] + v + original[3] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[3] + v ,
        original[16] + v + original[16] + v + original[10] + v + original[3] + v

    ]
    form_A_dict['SAD'] = [
        original[17] + v + original[9] + v ,  # Isolated
        original[17] + v + original[27] + v ,
        original[17] + v + original[31] + v ,
        original[17] + v + original[13] + v ,  # Isolated
        original[17] + v + original[26] + v + original[21] + v + original[27] + v,
        original[17] + v + original[26] + v + '\u0649' + v,

        original[17] + v + original[17] + v + original[13] + v , # Final + same front
        original[17] + v + original[17] + v + original[31] + v ,
        original[17] + v + original[17] + v + original[9] + v + original[9] + v ,
        original[17] + v + original[17] + v + original[27] + v + original[27] + v ,
        original[17] + v + original[17] + v + original[9] + v + original[31] + v ,

        original[17] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[9] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[27] + v + original[27] + v + original[3] + v  # Inital + alif end

    ]
    form_A_dict['DAD'] = [
        original[18] + v + original[8] + v ,  # Isolated
        original[18] + v + original[9] + v ,  # Isolated
        original[18] + v + original[10] + v ,  # Isolated
        original[18] + v + original[27] + v ,
        original[18] + v + original[31] + v ,
        original[18] + v + original[13] + v ,  # Isolated

        original[18] + v + original[18] + v + original[31] + v ,
        original[18] + v + original[18] + v + original[13] + v , # Final + same front
        original[18] + v + original[18] + v + original[9] + v + '\u0649' + v ,
        original[18] + v + original[18] + v + original[10] + v + original[27] + v ,
        original[18] + v + original[18] + v + original[9] + v + original[31] + v ,

        original[18] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[10] + v + original[27] + v + original[3] + v
    ]
    form_A_dict['TAH'] = [
        original[19] + v + original[9] + v ,  # Isolated
        original[19] + v + original[27] + v ,
        original[19] + v + original[31] + v ,

        original[19] + v + original[19] + v + original[31] + v ,
        original[19] + v + original[19] + v + original[27] + v + original[9] + v ,
        original[19] + v + original[19] + v + original[27] + v + original[31] + v ,

        original[19] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[19] + v + original[19] + v + original[27] + v + original[3] + v # medial 
    ]
    form_A_dict['ZAH'] = [
        original[20] + v + original[27] + v ,

        original[20] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[20] + v + original[20] + v + original[27] + v + original[3] + v # medial 
    ]
    form_A_dict['AIN'] = [
        original[21] + v + original[8] + v ,  # Isolated
        original[21] + v + original[27] + v ,
        original[21] + v + original[31] + v ,
        original[21] + v + original[26] + v + original[31] + v + original[29] + v ,  # Isolated

        original[21] + v + original[21] + v + original[31] + v ,
        original[21] + v + original[21] + v + original[8] + v + original[27] + v ,
        original[21] + v + original[21] + v + original[27] + v + original[27] + v ,
        original[21] + v + original[21] + v + original[27] + v + original[31] + v ,

        original[21] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[21] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[21] + v + original[27] + v + original[27] + v + original[3] + v,
        original[21] + v + original[8] + v + original[27] + v + original[3] + v

    ]
    form_A_dict['GHAIN'] = [
        original[22] + v + original[8] + v ,  # Isolated
        original[22] + v + original[27] + v ,
        original[22] + v + original[31] + v ,

        original[22] + v + original[22] + v + original[31] + v ,
        original[22] + v + original[22] + v + original[27] + v + original[27] + v ,
        original[22] + v + original[22] + v + original[27] + v + original[31] + v ,

        original[22] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[22] + v + original[27] + v + original[3] + v 
    ]
    form_A_dict['FEH'] = [
        original[23] + v + original[8] + v ,  # Isolated
        original[23] + v + original[9] + v ,  # Isolated
        original[23] + v + original[10] + v ,  # Isolated
        original[23] + v + original[27] + v ,
        original[23] + v + original[31] + v ,

        original[23] + v + original[23] + v + original[31] + v ,
        original[23] + v + original[23] + v + original[10] + v + original[27] + v ,
        original[23] + v + original[23] + v + original[27] + v + original[31] + v ,

        original[23] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[10] + v + original[27] + v + original[3] + v

    ]
    form_A_dict['QAF'] = [
        original[24] + v + original[9] + v ,  # Isolated
        original[24] + v + original[27] + v ,
        original[24] + v + original[31] + v ,

        original[24] + v + original[24] + v + original[31] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[9] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[27] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[31] + v ,

        original[24] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[24] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[24] + v + original[27] + v + original[9] + v + original[3] + v

    ]
    form_A_dict['KAF'] = [
        original[25] + v + original[1] + v ,  # Isolated
        original[25] + v + original[8] + v ,  # Isolated
        original[25] + v + original[9] + v ,  # Isolated
        original[25] + v + original[10] + v ,  # Isolated
        original[25] + v + original[26] + v ,  # Isolated
        original[25] + v + original[27] + v ,
        original[25] + v + original[31] + v ,
        
        original[25] + v + original[25] + v + original[1] + v ,
        original[25] + v + original[25] + v + original[26] + v, # Final + same front
        original[25] + v + original[25] + v + original[27] + v , # Final + same front
        original[25] + v + original[25] + v + original[31] + v ,
        original[25] + v + original[25] + v + original[27] + v + original[27] + v ,
        original[25] + v + original[25] + v + original[27] + v + original[31] + v ,

        original[25] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[26] + v + original[3] + v ,
        original[25] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[25] + v + original[25] + v + original[26] + v + original[3] + v , # medial + same front + alif end
        original[25] + v + original[25] + v + original[27] + v + original[3] + v 

    ]
    form_A_dict['LAM'] = [
        original[26] + v + original[8] + v ,  # Isolated
        original[26] + v + original[9] + v ,  # Isolated
        original[26] + v + original[10] + v ,  # Isolated
        original[26] + v + original[27] + v ,
        original[26] + v + original[31] + v ,

        original[26] + v + original[26] + v + original[27] + v , # Final + same front
        original[26] + v + original[26] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[8] + v ,
        original[26] + v + original[26] + v + original[10] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[27] + v + original[9] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[27] + v + original[31] + v ,

        original[26] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[29] + v + original[3] + v ,
        original[26] + v + original[8] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[26] + v + original[26] + v + original[27] + v + original[3] + v 

    ]
    form_A_dict['MEEM'] = [
        original[27] + v + original[8] + v ,  # Isolated
        original[27] + v + original[9] + v ,  # Isolated
        original[27] + v + original[10] + v ,  # Isolated
        original[27] + v + original[27] + v ,
        original[27] + v + original[31] + v ,
        original[27] + v + original[9] + v + original[27] + v + original[11] + v ,

        original[27] + v + original[27] + v + original[3] + v , # Final + same front
        original[27] + v + original[27] + v + original[27] + v ,
        original[27] + v + original[27] + v + original[9] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[27] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[10] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[8] + v + original[31] + v ,

        original[27] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[10] + v + original[3] + v , # Inital + alif end

    ]
    form_A_dict['NOON'] = [
        original[28] + v + original[8] + v ,  # Isolated
        original[28] + v + original[9] + v ,  # Isolated
        original[28] + v + original[10] + v ,  # Isolated
        original[28] + v + original[27] + v ,
        original[28] + v + original[31] + v ,

        original[28] + v + original[28] + v + original[13] + v , # Final + same front
        original[28] + v + original[28] + v + original[14] + v , # Final + same front
        original[28] + v + original[28] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[28] + v , # Final + same front
        original[28] + v + original[28] + v + original[31] + v , # Final + same front
        original[28] + v + original[28] + v + original[8] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[9] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[9] + v ,
        original[28] + v + original[28] + v + original[27] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[31] + v ,

        original[28] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[9] + v + original[3] + v ,

        original[28] + v + original[28] + v + original[27] + v + original[3] + v,   
        original[28] + v + original[28] + v + original[29] + v + original[3] + v   
    ]
    form_A_dict['HEH'] = [
        original[29] + v + original[8] + v ,  # Isolated
        original[29] + v + original[27] + v ,
        original[29] + v + original[31] + v ,

        original[29] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[29] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[29] + v + original[27] + v + original[8] + v + original[3] + v,
        original[29] + v + original[27] + v + original[27] + v + original[3] + v

    ]
    form_A_dict['YEH'] = [
        original[31] + v + original[8] + v ,  # Isolated
        original[31] + v + original[9] + v ,
        original[31] + v + original[10] + v ,
        original[31] + v + original[27] + v ,
        original[31] + v + original[31] + v ,
        
        original[31] + v + original[31] + v + original[13] + v , # Final + same front
        original[31] + v + original[31] + v + original[14] + v ,
        original[31] + v + original[31] + v + original[27] + v ,
        original[31] + v + original[31] + v + original[28] + v ,
        original[31] + v + original[31] + v + original[31] + v ,
        original[31] + v + original[31] + v + original[8] + v + original[31] + v ,
        original[31] + v + original[31] + v + original[27] + v + original[31] + v , # medial + same front + alif end
        original[31] + v + original[31] + v + original[27] + v + original[27] + v ,# medial + same front + alif end
        original[31] + v + original[31] + v + original[9] + v + original[31] + v ,
        
        original[31] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[31] + v + original[9] + v + original[3] + v ,
        original[31] + v + original[10] + v + original[3] + v ,
        original[31] + v + original[27] + v + original[3] + v ,
        original[31] + v + original[29] + v + original[3] + v ,
        original[31] + v + original[27] + v + original[27] + v + original[3] + v ,

        original[31] + v + original[31] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[31] + v + original[31] + v + original[29] + v + original[3] + v 

    ]
    return form_A_dict


def form_A_dict_one_by_one(tashkhil):
    if tashkhil:
        v = tashkhil_uni[3]
    else:
        v = ''
    # _____ FORM A DICT ONE BY ONE ______
    form_A_dict_one_by_one ={}
    form_A_dict_one_by_one['ALEEF'] = [
        original[25] + v + original[3] + v ,
        original[27] + v + original[3] + v ,
        original[3] + v + original[26] + v + original[26] + v + original[29] + v ,
        original[3] + v + original[25] + v + original[4] + v + original[13] + v,


        original[3] + v + original[25] + v + original[4] + v + original[13] + v,
        original[27] + v + original[27] + v + original[3] + v
    ]
    form_A_dict_one_by_one['BEH'] = [
        original[4] + v + original[8] + v ,  # Isolated
        original[4] + v + original[9] + v ,
        original[4] + v + original[10] + v ,
        original[4] + v + original[27] + v ,
        original[4] + v + original[31] + v ,
        original[4] + v + original[4] + v + original[31] + v , # Final + same front
        original[4] + v + original[4] + v + original[13] + v ,
        original[4] + v + original[4] + v + original[14] + v ,
        original[4] + v + original[4] + v + original[27] + v ,
        original[4] + v + original[4] + v + original[28] + v ,
        original[4] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[4] + v + original[27] + v + original[3] + v ,
        original[4] + v + original[29] + v + original[3] + v ,
        original[4] + v + original[8] + v + original[3] + v ,
        original[4] + v + original[10] + v + original[3] + v ,
        original[4] + v + original[4] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[4] + v + original[4] + v + original[29] + v + original[3] + v ,
        original[4] + v + original[4] + v + original[9] + v + original[31] + v , # final
        original[4] + v + original[4] + v + original[10] + v + original[31] + v , 
        original[3] + v + original[25] + v + original[4] + v + original[13] + v
    ]
    form_A_dict_one_by_one['TEH'] = [
        original[6] + v + original[8] + v ,  # Isolated
        original[6] + v + original[9] + v ,
        original[6] + v + original[10] + v ,
        original[6] + v + original[27] + v ,
        original[6] + v + original[31] + v ,
        
        original[6] + v + original[6] + v + original[13] + v , # Final + same front
        original[6] + v + original[6] + v + original[14] + v ,
        original[6] + v + original[6] + v + original[27] + v ,
        original[6] + v + original[6] + v + original[28] + v ,
        original[6] + v + original[6] + v + original[31] + v ,
        original[6] + v + original[6] + v + original[9] + v + original[8] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[8] + v + original[31] + v ,
        original[6] + v + original[6] + v + original[10] + v + original[31] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[27] + v + original[31] + v ,
        
        original[6] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[6] + v + original[9] + v + original[3] + v ,
        original[6] + v + original[10] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[3] + v ,
        original[6] + v + original[29] + v + original[3] + v ,
        original[6] + v + original[8] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[9] + v + original[8] + v + original[3] + v ,
        original[6] + v + original[9] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[10] + v + original[27] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[8] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[27] + v + original[9] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[10] + v + original[3] + v , # medial + same front + alif end
        
        original[6] + v + original[6] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[29] + v + original[3] + v 
    ]
    form_A_dict_one_by_one['THEH'] = [
        original[7] + v + original[8] + v ,  # Isolated
        original[7] + v + original[27] + v ,
        original[7] + v + original[31] + v ,
        
        original[7] + v + original[7] + v + original[13] + v , # Final + same front
        original[7] + v + original[7] + v + original[14] + v ,
        original[7] + v + original[7] + v + original[27] + v ,
        original[7] + v + original[7] + v + original[28] + v ,
        original[7] + v + original[7] + v + original[31] + v ,

        original[7] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[7] + v + original[7] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[7] + v + original[7] + v + original[29] + v + original[3] + v ,
    ]
    form_A_dict_one_by_one['JEEM'] = [
        original[8] + v + original[9] + v ,  # Isolated
        original[8] + v + original[27] + v ,
        original[8] + v + original[31] + v ,
        
        original[8] + v + original[8] + v + original[31] + v , # Final + same front
        original[8] + v + original[8] + v + original[27] + v + original[9] + v ,
        original[8] + v + original[8] + v + original[9] + v + original[31] + v,
        original[8] + v + original[8] + v + original[27] + v + original[31] + v,

        original[8] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[8] + v + original[8] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[8] + v + original[8] + v + original[29] + v + original[3] + v ,


        original[4] + v + original[8] + v ,  # Isolated
        original[4] + v + original[8] + v + original[3] + v ,
        original[6] + v + original[8] + v ,  # Isolated
        original[6] + v + original[6] + v + original[9] + v + original[8] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[8] + v + original[31] + v ,
        original[6] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[6] + v + original[8] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[9] + v + original[8] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[8] + v + original[3] + v , # medial + same front + alif end
        original[7] + v + original[8] + v ,  # Isolated
        original[9] + v + original[8] + v ,  # Isolated
        original[9] + v + original[9] + v + original[8] + v + original[31] + v,
        original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[10] + v + original[8] + v ,
        original[10] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[8] + v ,  # Isolated
        original[15] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[15] + v + original[8] + v + original[3] + v ,
        original[15] + v + original[27] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[8] + v ,  # Isolated
        original[16] + v + original[16] + v + original[8] + v , # Final + same front
        original[16] + v + original[16] + v + original[8] + v + original[31] + v ,
        original[16] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[16] + v + original[8] + v + original[3] + v ,
        original[18] + v + original[8] + v ,  # Isolated
        original[18] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[21] + v + original[8] + v ,  # Isolated
        original[21] + v + original[21] + v + original[8] + v + original[27] + v ,
        original[21] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[21] + v + original[8] + v + original[27] + v + original[3] + v,
        original[22] + v + original[8] + v ,  # Isolated
        original[22] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[8] + v ,  # Isolated
        original[23] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[8] + v ,  # Isolated
        original[25] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[8] + v ,  # Isolated
        original[26] + v + original[26] + v + original[8] + v + original[8] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[31] + v ,
        original[26] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[8] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v ,  # Isolated
        original[27] + v + original[27] + v + original[8] + v + original[31] + v ,
        original[27] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v ,  # Isolated
        original[28] + v + original[28] + v + original[8] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[9] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[31] + v ,
        original[28] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[9] + v + original[3] + v ,
        original[29] + v + original[8] + v ,  # Isolated
        original[29] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[29] + v + original[27] + v + original[8] + v + original[3] + v,
        original[31] + v + original[8] + v ,  # Isolated
        original[31] + v + original[31] + v + original[8] + v + original[31] + v ,
        original[31] + v + original[8] + v + original[3] + v 
        
    ]
    form_A_dict_one_by_one['HAH'] = [
        original[9] + v + original[8] + v ,  # Isolated
        original[9] + v + original[27] + v ,
        original[9] + v + original[31] + v ,
        
        original[9] + v + original[9] + v + original[31] + v , # Final + same front
        original[9] + v + original[9] + v + original[27] + v + original[31] + v ,
        original[9] + v + original[9] + v + original[8] + v + original[31] + v,

        original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[9] + v + original[27] + v + original[3] + v,   # Inital + alif end


        original[4] + v + original[9] + v ,
        original[4] + v + original[4] + v + original[9] + v + original[31] + v , # final
        original[6] + v + original[9] + v ,
        original[6] + v + original[6] + v + original[9] + v + original[8] + v , # medial + same front + alif end
        original[6] + v + original[9] + v + original[3] + v ,
        original[6] + v + original[9] + v + original[8] + v + original[3] + v ,
        original[6] + v + original[9] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[27] + v + original[9] + v + original[3] + v ,
        original[8] + v + original[9] + v ,  # Isolated
        original[8] + v + original[8] + v + original[27] + v + original[9] + v ,
        original[8] + v + original[8] + v + original[9] + v + original[31] + v,
        original[10] + v + original[9] + v ,  # Isolated
        original[15] + v + original[9] + v ,  # Isolated
        original[15] + v + original[15] + v + original[27] + v + original[9] + v ,
        original[15] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[15] + v + original[9] + v + original[3] + v ,
        original[16] + v + original[9] + v ,  # Isolated
        original[16] + v + original[16] + v + original[9] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[31] + v ,
        original[16] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[16] + v + original[9] + v + original[3] + v ,
        original[17] + v + original[9] + v ,  # Isolated
        original[17] + v + original[17] + v + original[9] + v + original[9] + v ,
        original[17] + v + original[17] + v + original[9] + v + original[31] + v ,
        original[17] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[9] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[9] + v ,  # Isolated
        original[18] + v + original[18] + v + original[9] + v + '\u0649' + v ,
        original[18] + v + original[18] + v + original[9] + v + original[31] + v ,
        original[18] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[9] + v ,  # Isolated
        original[19] + v + original[19] + v + original[27] + v + original[9] + v ,
        original[19] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[9] + v ,  # Isolated
        original[23] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[24] + v + original[9] + v ,  # Isolated
        original[24] + v + original[24] + v + original[27] + v + original[9] + v ,
        original[24] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[24] + v + original[27] + v + original[9] + v + original[3] + v,
        original[25] + v + original[9] + v ,  # Isolated
        original[25] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[9] + v ,  # Isolated
        original[26] + v + original[26] + v + original[27] + v + original[9] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[31] + v ,
        original[26] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v ,  # Isolated
        original[27] + v + original[9] + v + original[27] + v + original[11] + v ,
        original[27] + v + original[27] + v + original[9] + v + original[31] + v ,
        original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[9] + v ,  # Isolated
        original[28] + v + original[28] + v + original[9] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[9] + v ,
        original[28] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[9] + v + original[3] + v ,
        original[31] + v + original[9] + v ,
        original[31] + v + original[31] + v + original[9] + v + original[31] + v ,
        original[31] + v + original[9] + v + original[3] + v ,
        
    ]
    form_A_dict_one_by_one['KHAH'] = [
        original[10] + v + original[9] + v ,  # Isolated
        original[10] + v + original[27] + v ,
        original[10] + v + original[8] + v ,
        original[10] + v + original[31] + v ,

        original[10] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[10] + v + original[31] + v + original[3] + v,  # Inital + alif end


        original[4] + v + original[10] + v ,
        original[4] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[4] + v + original[10] + v + original[3] + v ,
        original[4] + v + original[4] + v + original[10] + v + original[31] + v , 
        original[6] + v + original[10] + v ,
        original[6] + v + original[6] + v + original[10] + v + original[31] + v , # medial + same front + alif end
        original[6] + v + original[10] + v + original[3] + v ,
        original[6] + v + original[10] + v + original[27] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[10] + v + original[3] + v , # medial + same front + alif end
        original[15] + v + original[10] + v ,  # Isolated
        original[15] + v + original[15] + v + original[10] + v + original[31] + v ,
        original[15] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[15] + v + original[10] + v + original[3] + v,
        original[16] + v + original[10] + v ,  # Isolated
        original[16] + v + original[16] + v + original[10] + v, # Final + same front
        original[16] + v + original[16] + v + original[27] + v + original[10] + v ,
        original[16] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[16] + v + original[10] + v + original[3] + v,
        original[17] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[10] + v ,  # Isolated
        original[18] + v + original[18] + v + original[10] + v + original[27] + v ,
        original[18] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[10] + v ,  # Isolated
        original[23] + v + original[23] + v + original[10] + v + original[27] + v ,
        original[23] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[10] + v + original[27] + v + original[3] + v,
        original[25] + v + original[10] + v ,  # Isolated
        original[25] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[10] + v ,  # Isolated
        original[26] + v + original[26] + v + original[10] + v + original[27] + v ,
        original[26] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v ,  # Isolated
        original[27] + v + original[27] + v + original[10] + v + original[31] + v ,
        original[27] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[10] + v ,  # Isolated
        original[28] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[31] + v + original[10] + v ,
        original[31] + v + original[10] + v + original[3] + v
        
    ]
    form_A_dict_one_by_one['SEEN'] = [
        original[15] + v + original[8] + v ,  # Isolated
        original[15] + v + original[9] + v ,  # Isolated
        original[15] + v + original[10] + v ,  # Isolated
        original[15] + v + original[13] + v ,  # Isolated
        original[15] + v + original[27] + v ,
        original[15] + v + original[31] + v ,
        
        original[15] + v + original[15] + v + original[13] + v , # Final + same front
        original[15] + v + original[15] + v + original[31] + v ,
        original[15] + v + original[15] + v + original[27] + v + original[9] + v ,
        original[15] + v + original[15] + v + original[27] + v + original[27] + v ,
        original[15] + v + original[15] + v + original[10] + v + original[31] + v ,

        original[15] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[15] + v + original[15] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[15] + v + original[15] + v + original[29] + v + original[3] + v ,
        original[15] + v + original[15] + v + original[8] + v + original[3] + v ,
        original[15] + v + original[15] + v + original[9] + v + original[3] + v ,
        original[15] + v + original[15] + v + original[10] + v + original[3] + v

    ]
    form_A_dict_one_by_one['SHEEN'] = [
        original[16] + v + original[8] + v ,  # Isolated
        original[16] + v + original[9] + v ,  # Isolated
        original[16] + v + original[10] + v ,  # Isolated
        original[16] + v + original[13] + v ,  # Isolated
        original[16] + v + original[27] + v ,
        original[16] + v + original[31] + v ,
        
        original[16] + v + original[16] + v + original[8] + v , # Final + same front
        original[16] + v + original[16] + v + original[9] + v ,
        original[16] + v + original[16] + v + original[10] + v, # Final + same front
        original[16] + v + original[16] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[13] + v , # Final + same front
        original[16] + v + original[16] + v + original[31] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[8] + v + original[31] + v ,
        original[16] + v + original[16] + v + original[27] + v + original[10] + v ,
        original[16] + v + original[16] + v + original[27] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[31] + v ,

        original[16] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[16] + v + original[16] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[16] + v + original[16] + v + original[29] + v + original[3] + v ,
        original[16] + v + original[16] + v + original[8] + v + original[3] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[3] + v ,
        original[16] + v + original[16] + v + original[10] + v + original[3] + v

    ]
    form_A_dict_one_by_one['SAD'] = [
        original[17] + v + original[9] + v ,  # Isolated
        original[17] + v + original[27] + v ,
        original[17] + v + original[31] + v ,
        original[17] + v + original[13] + v ,  # Isolated
        original[17] + v + original[26] + v + original[21] + v + original[27] + v,
        original[17] + v + original[26] + v + '\u0649' + v,

        original[17] + v + original[17] + v + original[13] + v , # Final + same front
        original[17] + v + original[17] + v + original[31] + v ,
        original[17] + v + original[17] + v + original[9] + v + original[9] + v ,
        original[17] + v + original[17] + v + original[27] + v + original[27] + v ,
        original[17] + v + original[17] + v + original[9] + v + original[31] + v ,

        original[17] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[9] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[27] + v + original[27] + v + original[3] + v  # Inital + alif end

    ]
    form_A_dict_one_by_one['DAD'] = [
        original[18] + v + original[8] + v ,  # Isolated
        original[18] + v + original[9] + v ,  # Isolated
        original[18] + v + original[10] + v ,  # Isolated
        original[18] + v + original[27] + v ,
        original[18] + v + original[31] + v ,
        original[18] + v + original[13] + v ,  # Isolated

        original[18] + v + original[18] + v + original[31] + v ,
        original[18] + v + original[18] + v + original[13] + v , # Final + same front
        original[18] + v + original[18] + v + original[9] + v + '\u0649' + v ,
        original[18] + v + original[18] + v + original[10] + v + original[27] + v ,
        original[18] + v + original[18] + v + original[9] + v + original[31] + v ,

        original[18] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[10] + v + original[27] + v + original[3] + v
    ]
    form_A_dict_one_by_one['TAH'] = [
        original[19] + v + original[9] + v ,  # Isolated
        original[19] + v + original[27] + v ,
        original[19] + v + original[31] + v ,

        original[19] + v + original[19] + v + original[31] + v ,
        original[19] + v + original[19] + v + original[27] + v + original[9] + v ,
        original[19] + v + original[19] + v + original[27] + v + original[31] + v ,

        original[19] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[19] + v + original[19] + v + original[27] + v + original[3] + v # medial 
    ]
    form_A_dict_one_by_one['ZAH'] = [
        original[20] + v + original[27] + v ,

        original[20] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[20] + v + original[20] + v + original[27] + v + original[3] + v # medial 
    ]
    form_A_dict_one_by_one['AIN'] = [
        original[21] + v + original[8] + v ,  # Isolated
        original[21] + v + original[27] + v ,
        original[21] + v + original[31] + v ,
        original[21] + v + original[26] + v + original[31] + v + original[29] + v ,  # Isolated

        original[21] + v + original[21] + v + original[31] + v ,
        original[21] + v + original[21] + v + original[8] + v + original[27] + v ,
        original[21] + v + original[21] + v + original[27] + v + original[27] + v ,
        original[21] + v + original[21] + v + original[27] + v + original[31] + v ,

        original[21] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[21] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[21] + v + original[27] + v + original[27] + v + original[3] + v,
        original[21] + v + original[8] + v + original[27] + v + original[3] + v

    ]
    form_A_dict_one_by_one['GHAIN'] = [
        original[22] + v + original[8] + v ,  # Isolated
        original[22] + v + original[27] + v ,
        original[22] + v + original[31] + v ,

        original[22] + v + original[22] + v + original[31] + v ,
        original[22] + v + original[22] + v + original[27] + v + original[27] + v ,
        original[22] + v + original[22] + v + original[27] + v + original[31] + v ,

        original[22] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[22] + v + original[27] + v + original[3] + v 
    ]
    form_A_dict_one_by_one['FEH'] = [
        original[23] + v + original[8] + v ,  # Isolated
        original[23] + v + original[9] + v ,  # Isolated
        original[23] + v + original[10] + v ,  # Isolated
        original[23] + v + original[27] + v ,
        original[23] + v + original[31] + v ,

        original[23] + v + original[23] + v + original[31] + v ,
        original[23] + v + original[23] + v + original[10] + v + original[27] + v ,
        original[23] + v + original[23] + v + original[27] + v + original[31] + v ,

        original[23] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[10] + v + original[27] + v + original[3] + v

    ]
    form_A_dict_one_by_one['QAF'] = [
        original[24] + v + original[9] + v ,  # Isolated
        original[24] + v + original[27] + v ,
        original[24] + v + original[31] + v ,

        original[24] + v + original[24] + v + original[31] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[9] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[27] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[31] + v ,

        original[24] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[24] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[24] + v + original[27] + v + original[9] + v + original[3] + v

    ]
    form_A_dict_one_by_one['KAF'] = [
        original[25] + v + original[1] + v ,  # Isolated
        original[25] + v + original[8] + v ,  # Isolated
        original[25] + v + original[9] + v ,  # Isolated
        original[25] + v + original[10] + v ,  # Isolated
        original[25] + v + original[26] + v ,  # Isolated
        original[25] + v + original[27] + v ,
        original[25] + v + original[31] + v ,
        
        original[25] + v + original[25] + v + original[1] + v ,
        original[25] + v + original[25] + v + original[26] + v, # Final + same front
        original[25] + v + original[25] + v + original[27] + v , # Final + same front
        original[25] + v + original[25] + v + original[31] + v ,
        original[25] + v + original[25] + v + original[27] + v + original[27] + v ,
        original[25] + v + original[25] + v + original[27] + v + original[31] + v ,

        original[25] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[26] + v + original[3] + v ,
        original[25] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[25] + v + original[25] + v + original[26] + v + original[3] + v , # medial + same front + alif end
        original[25] + v + original[25] + v + original[27] + v + original[3] + v ,

        original[3] + v + original[25] + v + original[4] + v + original[13] + v,
        original[25] + v + original[3] + v 

    ]
    form_A_dict_one_by_one['LAM'] = [
        original[26] + v + original[8] + v ,  # Isolated
        original[26] + v + original[9] + v ,  # Isolated
        original[26] + v + original[10] + v ,  # Isolated
        original[26] + v + original[27] + v ,
        original[26] + v + original[31] + v ,

        original[26] + v + original[26] + v + original[27] + v , # Final + same front
        original[26] + v + original[26] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[8] + v ,
        original[26] + v + original[26] + v + original[10] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[27] + v + original[9] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[27] + v + original[31] + v ,

        original[26] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[29] + v + original[3] + v ,
        original[26] + v + original[8] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end

        original[26] + v + original[26] + v + original[27] + v + original[3] + v ,

        original[3] + v + original[26] + v + original[26] + v + original[29] + v ,
        original[17] + v + original[26] + v + original[21] + v + original[27] + v,
        original[17] + v + original[26] + v + '\u0649' + v,
        original[21] + v + original[26] + v + original[31] + v + original[29] + v ,  # Isolated
        original[25] + v + original[26] + v ,  # Isolated
        original[25] + v + original[25] + v + original[26] + v, # Final + same front
        original[25] + v + original[26] + v + original[3] + v ,
        original[25] + v + original[25] + v + original[26] + v + original[3] + v # medial + same front + alif end
        
    ]
    form_A_dict_one_by_one['MEEM'] = [
        original[27] + v + original[8] + v ,  # Isolated
        original[27] + v + original[9] + v ,  # Isolated
        original[27] + v + original[10] + v ,  # Isolated
        original[27] + v + original[27] + v ,
        original[27] + v + original[31] + v ,
        original[27] + v + original[9] + v + original[27] + v + original[11] + v ,

        original[27] + v + original[27] + v + original[3] + v , # Final + same front
        original[27] + v + original[27] + v + original[27] + v ,
        original[27] + v + original[27] + v + original[9] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[27] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[10] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[8] + v + original[31] + v ,

        original[27] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[27] + v + original[8] + v + original[10] + v + original[3] + v , # Inital + alif end


        original[27] + v + original[3] + v ,
        original[4] + v + original[27] + v ,
        original[4] + v + original[4] + v + original[27] + v ,
        original[4] + v + original[27] + v + original[3] + v ,
        original[4] + v + original[4] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[27] + v ,
        original[6] + v + original[6] + v + original[27] + v ,
        original[6] + v + original[6] + v + original[27] + v + original[31] + v ,
        original[6] + v + original[27] + v + original[3] + v ,
        original[6] + v + original[8] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[9] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[10] + v + original[27] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[8] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[27] + v + original[9] + v + original[3] + v ,
        original[6] + v + original[27] + v + original[10] + v + original[3] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[7] + v + original[27] + v ,
        original[7] + v + original[7] + v + original[27] + v ,
        original[7] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[7] + v + original[7] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[8] + v + original[27] + v ,
        original[8] + v + original[8] + v + original[27] + v + original[9] + v ,
        original[8] + v + original[8] + v + original[27] + v + original[31] + v,
        original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[8] + v + original[8] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[9] + v + original[27] + v ,
        original[9] + v + original[9] + v + original[27] + v + original[31] + v ,
        original[9] + v + original[27] + v + original[3] + v, # Inital + alif end
        original[15] + v + original[27] + v ,
        original[15] + v + original[15] + v + original[27] + v + original[9] + v ,
        original[15] + v + original[15] + v + original[27] + v + original[27] + v ,
        original[15] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[15] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[16] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[27] + v ,
        original[16] + v + original[16] + v + original[27] + v + original[10] + v ,
        original[16] + v + original[16] + v + original[27] + v + original[27] + v ,
        original[16] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[16] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[17] + v + original[27] + v ,
        original[17] + v + original[26] + v + original[21] + v + original[27] + v,
        original[17] + v + original[17] + v + original[27] + v + original[27] + v ,
        original[17] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[17] + v + original[27] + v + original[27] + v + original[3] + v,  # Inital + alif end
        original[18] + v + original[27] + v ,
        original[18] + v + original[18] + v + original[10] + v + original[27] + v ,
        original[18] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[18] + v + original[10] + v + original[27] + v + original[3] + v,
        original[19] + v + original[27] + v ,
        original[19] + v + original[19] + v + original[27] + v + original[9] + v ,
        original[19] + v + original[19] + v + original[27] + v + original[31] + v ,
        original[19] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[19] + v + original[19] + v + original[27] + v + original[3] + v, # medial 
        original[20] + v + original[27] + v ,
        original[20] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[20] + v + original[20] + v + original[27] + v + original[3] + v, # medial 
        original[21] + v + original[27] + v ,
        original[21] + v + original[21] + v + original[8] + v + original[27] + v ,
        original[21] + v + original[21] + v + original[27] + v + original[27] + v ,
        original[21] + v + original[21] + v + original[27] + v + original[31] + v ,
        original[21] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[21] + v + original[27] + v + original[27] + v + original[3] + v,
        original[21] + v + original[8] + v + original[27] + v + original[3] + v,
        original[22] + v + original[27] + v ,
        original[22] + v + original[22] + v + original[27] + v + original[27] + v ,
        original[22] + v + original[22] + v + original[27] + v + original[31] + v ,
        original[22] + v + original[27] + v + original[3] + v,
        original[23] + v + original[27] + v ,
        original[23] + v + original[23] + v + original[10] + v + original[27] + v ,
        original[23] + v + original[23] + v + original[27] + v + original[31] + v ,
        original[23] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[23] + v + original[10] + v + original[27] + v + original[3] + v,
        original[24] + v + original[27] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[9] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[27] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[31] + v ,
        original[24] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[24] + v + original[27] + v + original[9] + v + original[3] + v,
        original[25] + v + original[27] + v ,
        original[25] + v + original[25] + v + original[27] + v , # Final + same front
        original[25] + v + original[25] + v + original[27] + v + original[27] + v ,
        original[25] + v + original[25] + v + original[27] + v + original[31] + v ,
        original[25] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[27] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[25] + v + original[25] + v + original[27] + v + original[3] + v ,
        original[26] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[27] + v , # Final + same front
        original[26] + v + original[26] + v + original[10] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[27] + v + original[9] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[27] + v ,
        original[26] + v + original[26] + v + original[27] + v + original[31] + v ,
        original[26] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[10] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[27] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[26] + v + original[26] + v + original[27] + v + original[3] + v ,
        original[28] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[27] + v + original[31] + v ,
        original[28] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[28] + v + original[27] + v + original[3] + v,   
        original[29] + v + original[27] + v ,
        original[29] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[29] + v + original[27] + v + original[8] + v + original[3] + v,
        original[29] + v + original[27] + v + original[27] + v + original[3] + v,
        original[31] + v + original[27] + v ,
        original[31] + v + original[31] + v + original[27] + v ,
        original[31] + v + original[31] + v + original[27] + v + original[31] + v , # medial + same front + alif end
        original[31] + v + original[31] + v + original[27] + v + original[27] + v ,# medial + same front + alif end
        original[31] + v + original[27] + v + original[3] + v ,
        original[31] + v + original[27] + v + original[27] + v + original[3] + v ,
        original[31] + v + original[31] + v + original[27] + v + original[3] + v # medial + same front + alif end
    ]
    form_A_dict_one_by_one['NOON'] = [
        original[28] + v + original[8] + v ,  # Isolated
        original[28] + v + original[9] + v ,  # Isolated
        original[28] + v + original[10] + v ,  # Isolated
        original[28] + v + original[27] + v ,
        original[28] + v + original[31] + v ,

        original[28] + v + original[28] + v + original[13] + v , # Final + same front
        original[28] + v + original[28] + v + original[14] + v , # Final + same front
        original[28] + v + original[28] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[28] + v , # Final + same front
        original[28] + v + original[28] + v + original[31] + v , # Final + same front
        original[28] + v + original[28] + v + original[8] + v + original[27] + v ,
        original[28] + v + original[28] + v + original[9] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[9] + v ,
        original[28] + v + original[28] + v + original[27] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[31] + v ,

        original[28] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[9] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[10] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[9] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[8] + v + original[9] + v + original[3] + v ,

        original[28] + v + original[28] + v + original[27] + v + original[3] + v,   
        original[28] + v + original[28] + v + original[29] + v + original[3] + v,

        original[4] + v + original[4] + v + original[28] + v ,
        original[6] + v + original[6] + v + original[28] + v ,
        original[7] + v + original[7] + v + original[28] + v ,
        original[31] + v + original[31] + v + original[28] + v 
    ]
    form_A_dict_one_by_one['HEH'] = [
        original[29] + v + original[8] + v ,  # Isolated
        original[29] + v + original[27] + v ,
        original[29] + v + original[31] + v ,

        original[29] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[29] + v + original[27] + v + original[3] + v , # Inital + alif end
        original[29] + v + original[27] + v + original[8] + v + original[3] + v,
        original[29] + v + original[27] + v + original[27] + v + original[3] + v,


        original[3] + v + original[26] + v + original[26] + v + original[29] + v ,
        original[4] + v + original[29] + v + original[3] + v ,
        original[4] + v + original[4] + v + original[29] + v + original[3] + v ,
        original[6] + v + original[29] + v + original[3] + v ,
        original[6] + v + original[6] + v + original[29] + v + original[3] + v ,
        original[7] + v + original[7] + v + original[29] + v + original[3] + v ,
        original[8] + v + original[8] + v + original[29] + v + original[3] + v,
        original[15] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[15] + v + original[15] + v + original[29] + v + original[3] + v ,
        original[16] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[16] + v + original[16] + v + original[29] + v + original[3] + v ,
        original[21] + v + original[26] + v + original[31] + v + original[29] + v ,  # Isolated
        original[26] + v + original[29] + v + original[3] + v ,
        original[28] + v + original[29] + v + original[3] + v , # Inital + alif end
        original[28] + v + original[28] + v + original[29] + v + original[3] + v,
        original[31] + v + original[29] + v + original[3] + v ,
        original[31] + v + original[31] + v + original[29] + v + original[3] + v 

    ]
    form_A_dict_one_by_one['YEH'] = [
        original[31] + v + original[8] + v ,  # Isolated
        original[31] + v + original[9] + v ,
        original[31] + v + original[10] + v ,
        original[31] + v + original[27] + v ,
        original[31] + v + original[31] + v ,
        
        original[31] + v + original[31] + v + original[13] + v , # Final + same front
        original[31] + v + original[31] + v + original[14] + v ,
        original[31] + v + original[31] + v + original[27] + v ,
        original[31] + v + original[31] + v + original[28] + v ,
        original[31] + v + original[31] + v + original[31] + v ,
        original[31] + v + original[31] + v + original[8] + v + original[31] + v ,
        original[31] + v + original[31] + v + original[27] + v + original[31] + v , # medial + same front + alif end
        original[31] + v + original[31] + v + original[27] + v + original[27] + v ,# medial + same front + alif end
        original[31] + v + original[31] + v + original[9] + v + original[31] + v ,
        
        original[31] + v + original[8] + v + original[3] + v , # Inital + alif end
        original[31] + v + original[9] + v + original[3] + v ,
        original[31] + v + original[10] + v + original[3] + v ,
        original[31] + v + original[27] + v + original[3] + v ,
        original[31] + v + original[29] + v + original[3] + v ,
        original[31] + v + original[27] + v + original[27] + v + original[3] + v ,

        original[31] + v + original[31] + v + original[27] + v + original[3] + v , # medial + same front + alif end
        original[31] + v + original[31] + v + original[29] + v + original[3] + v,

        original[4] + v + original[31] + v ,
        original[4] + v + original[4] + v + original[31] + v , # Final + same front
        original[4] + v + original[4] + v + original[9] + v + original[31] + v , # final
        original[4] + v + original[4] + v + original[10] + v + original[31] + v , 
        original[6] + v + original[31] + v ,
        original[6] + v + original[6] + v + original[31] + v ,
        original[6] + v + original[6] + v + original[8] + v + original[31] + v ,
        original[6] + v + original[6] + v + original[10] + v + original[31] + v , # medial + same front + alif end
        original[6] + v + original[6] + v + original[27] + v + original[31] + v ,
        original[7] + v + original[31] + v ,
        original[7] + v + original[7] + v + original[31] + v ,
        original[8] + v + original[31] + v ,
        original[8] + v + original[8] + v + original[31] + v , # Final + same front
        original[8] + v + original[8] + v + original[9] + v + original[31] + v,
        original[8] + v + original[8] + v + original[27] + v + original[31] + v,
        original[9] + v + original[31] + v ,
        original[9] + v + original[9] + v + original[31] + v , # Final + same front
        original[9] + v + original[9] + v + original[27] + v + original[31] + v ,
        original[9] + v + original[9] + v + original[8] + v + original[31] + v,
        original[10] + v + original[31] + v ,
        original[10] + v + original[31] + v + original[3] + v,  # Inital + alif end
        original[15] + v + original[31] + v ,
        original[15] + v + original[15] + v + original[31] + v ,
        original[15] + v + original[15] + v + original[10] + v + original[31] + v ,
        original[16] + v + original[31] + v ,
        original[16] + v + original[16] + v + original[31] + v ,
        original[16] + v + original[16] + v + original[8] + v + original[31] + v ,
        original[16] + v + original[16] + v + original[9] + v + original[31] + v ,
        original[17] + v + original[31] + v ,
        original[17] + v + original[17] + v + original[31] + v ,
        original[17] + v + original[17] + v + original[9] + v + original[31] + v ,
        original[18] + v + original[31] + v ,
        original[18] + v + original[18] + v + original[31] + v ,
        original[18] + v + original[18] + v + original[9] + v + original[31] + v ,
        original[19] + v + original[31] + v ,
        original[19] + v + original[19] + v + original[31] + v ,
        original[19] + v + original[19] + v + original[27] + v + original[31] + v ,
        original[21] + v + original[31] + v ,
        original[21] + v + original[26] + v + original[31] + v + original[29] + v ,  # Isolated
        original[21] + v + original[21] + v + original[31] + v ,
        original[21] + v + original[21] + v + original[27] + v + original[31] + v ,
        original[22] + v + original[31] + v ,
        original[22] + v + original[22] + v + original[31] + v ,
        original[22] + v + original[22] + v + original[27] + v + original[31] + v ,
        original[23] + v + original[31] + v ,
        original[23] + v + original[23] + v + original[31] + v ,
        original[23] + v + original[23] + v + original[27] + v + original[31] + v ,
        original[24] + v + original[31] + v ,
        original[24] + v + original[24] + v + original[31] + v ,
        original[24] + v + original[24] + v + original[27] + v + original[31] + v ,
        original[25] + v + original[31] + v ,
        original[25] + v + original[25] + v + original[31] + v ,
        original[25] + v + original[25] + v + original[27] + v + original[31] + v ,
        original[26] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[9] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[8] + v + original[31] + v ,
        original[26] + v + original[26] + v + original[27] + v + original[31] + v ,
        original[27] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[9] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[27] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[10] + v + original[31] + v ,
        original[27] + v + original[27] + v + original[8] + v + original[31] + v ,
        original[28] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[31] + v , # Final + same front
        original[28] + v + original[28] + v + original[9] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[27] + v + original[31] + v ,
        original[28] + v + original[28] + v + original[8] + v + original[31] + v ,
        original[29] + v + original[31] + v 
    ]
    return form_A_dict_one_by_one


#_______configuration______________
# font_name = 'me_quran'
font_name = font_arg
# size = 48
size = size_arg
font_size = Pt(size)
space = '          '
if repetition_arg != '':
    rep = repetition_arg
else:
    rep = 6
if dot_arg != '':
    dot = dot_arg
else:
    dot = 10
space = ''
for x in range(dot):
    space += ' ' 
#_________________________________
from docx.shared import RGBColor

document = Document()
styles = document.styles
style = styles.add_style('Test', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = font_name
style.font.size = font_size
text = ''
for x in original:
    text += x + space
# print(text)
sum_text_b = ''
if repetition_arg != '':
    count = 0
    cnt = 0
    for x in forms_b:
        count += 1
        cnt += 1
        sum_text_b += x + space
        if count == rep or cnt == len(forms_b):
            p = document.add_paragraph(sum_text_b, style='Test')
            sum_text_b = ''
            count = 0
else:
    count = 0
    cnt = 0
    para = 0
    rep = [7, 4, 6, 4, 4, 4, 4,
           2, 2, 2, 2, 4, 4, 4,
           4, 4, 4, 4, 4, 4, 4,
           4, 6, 4, 4, 4, 2, 4]
    for x in forms_b:
        count += 1
        cnt += 1
        sum_text_b += x + space
        if count == rep[para] or cnt == len(forms_b):
            p = document.add_paragraph(sum_text_b, style='Test')
            sum_text_b = ''
            count = 0
            para += 1

document1 = Document()
styles = document1.styles
style = styles.add_style('Test', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = font_name
style.font.size = font_size
form_A_dict = form_A_dict_one_by_one(tashkhil=True)
for key in form_A_dict:
    sum_char = ''
    p1 = document1.add_paragraph(key, style='Test')
    for x in form_A_dict[key]:
        sum_char += x + space
    p2 = document1.add_paragraph(sum_char, style='Test')

document2 = Document()
styles = document2.styles
style = styles.add_style('Test', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = font_name
style.font.size = font_size
form_A_dict_one_by_one = form_A_dict_one_by_one(tashkhil=False)
for key in form_A_dict_one_by_one:
    sum_char = ''
    p1 = document2.add_paragraph(key, style='Test')
    for x in form_A_dict_one_by_one[key]:
        sum_char += x + space
    p2 = document2.add_paragraph(sum_char, style='Test')

# run = document.add_heading().add_run("TITLE")
# font = run.font
# font.color.rgb = RGBColor(255,255,255)
# run = document1.add_heading().add_run("TITLE")
# font = run.font
# font.color.rgb = RGBColor(255,255,255)
# run = document2.add_heading().add_run("TITLE")
# font = run.font
# font.color.rgb = RGBColor(255,255,255)

document.save('Forms B^'+font_name+'_'+str(size)+'.docx')
document1.save('Forms A Full_Type_one_by_one(with tashkhil)^'+font_name+'_'+str(size)+'.docx')
document2.save('Forms A Full_Type_one_by_one^'+font_name+'_'+str(size)+'.docx')
